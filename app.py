# Import necessary libraries
from flask import Flask, request, render_template, send_file, jsonify
import os
import certifi
import ssl
import urllib.request
import whisper
import whisperx
from huggingface_hub import login
from pyannote.audio import Pipeline
import pandas as pd
from transformers import pipeline, AutoTokenizer
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import time
import uuid
import subprocess
import math
import pytz
import logging
import re

# Configure logging with UTF-8 encoding
logging.basicConfig(level=logging.DEBUG, encoding='utf-8')
logger = logging.getLogger(__name__)

# Initialize Flask app and set up upload/output directories
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Set up SSL context for secure HTTPS requests
cert_path = certifi.where()
context = ssl.create_default_context(cafile=cert_path)
opener = urllib.request.build_opener(urllib.request.HTTPSHandler(context=context))
urllib.request.install_opener(opener)

# Authenticate with Hugging Face using environment token
HF_TOKEN = os.getenv("HF_TOKEN")
login(token=HF_TOKEN)

# Format seconds into human-readable time (hours, minutes, seconds)
def format_processing_time(seconds):
    """Convert seconds to a human-readable format (hours, minutes, seconds)."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    remaining_seconds = seconds % 60
    parts = []
    if hours > 0:
        parts.append(f"{hours} hour{'s' if hours != 1 else ''}")
    if minutes > 0:
        parts.append(f"{minutes} minute{'s' if minutes != 1 else ''}")
    if remaining_seconds > 0 or not parts:
        parts.append(f"{remaining_seconds:.2f} second{'s' if remaining_seconds != 1 else ''}")
    return " ".join(parts)

# Convert seconds to [MM:SS] format for timestamps
def format_timestamp(seconds):
    """Convert seconds to [MM:SS] format for internal use."""
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"[{minutes}:{secs:02d}]"

# Convert seconds to MM:SS.mmm format with milliseconds
def format_timestamp_ms(seconds):
    """Convert seconds to MM:SS.mmm format for internal use (with milliseconds)."""
    minutes = int(seconds // 60)
    secs = seconds % 60
    return f"{minutes:02d}:{secs:06.3f}"

# Split text into chunks for summarization, preserving sentences
def split_text_into_chunks(text, tokenizer, max_tokens=1000, overlap_tokens=300):
    """Split text into chunks with specified token counts and overlaps, preserving full sentences."""
    tokens = tokenizer(text, return_tensors="pt", truncation=False).input_ids[0]
    total_tokens = len(tokens)
    logger.debug(f"Total tokens for text: {total_tokens}")
    chunks = []
    chunk_token_counts = []
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() + ('.' if s and not s.endswith('.') else '') for s in sentences if s.strip()]
    
    current_chunk = []
    current_token_count = 0
    sentence_index = 0
    
    while sentence_index < len(sentences):
        # Create first chunk (~1000 tokens)
        if not chunks:
            while sentence_index < len(sentences) and current_token_count < max_tokens:
                sentence = sentences[sentence_index]
                sentence_tokens = len(tokenizer(sentence, return_tensors="pt").input_ids[0])
                if current_token_count + sentence_tokens <= max_tokens or not current_chunk:
                    current_chunk.append(sentence)
                    current_token_count += sentence_tokens
                    sentence_index += 1
                else:
                    break
            chunks.append(' '.join(current_chunk))
            chunk_token_counts.append(current_token_count)
            logger.debug(f"Chunk 1 created: {current_token_count} tokens")
        
        # Create subsequent chunks with overlap (~300 tokens from previous + ~700 new)
        else:
            # Get last ~300 tokens from previous chunk
            prev_chunk = chunks[-1]
            prev_sentences = re.split(r'(?<=[.!?])\s+', prev_chunk)
            prev_sentences = [s.strip() for s in prev_sentences if s.strip()]
            overlap_chunk = []
            overlap_token_count = 0
            for sentence in reversed(prev_sentences):
                sentence_tokens = len(tokenizer(sentence, return_tensors="pt").input_ids[0])
                if overlap_token_count + sentence_tokens <= overlap_tokens or not overlap_chunk:
                    overlap_chunk.insert(0, sentence)
                    overlap_token_count += sentence_tokens
                else:
                    break
            
            # Add new ~700 tokens
            new_chunk = overlap_chunk.copy()
            new_token_count = overlap_token_count
            while sentence_index < len(sentences) and new_token_count < max_tokens:
                sentence = sentences[sentence_index]
                sentence_tokens = len(tokenizer(sentence, return_tensors="pt").input_ids[0])
                if new_token_count + sentence_tokens <= max_tokens or not new_chunk:
                    new_chunk.append(sentence)
                    new_token_count += sentence_tokens
                    sentence_index += 1
                else:
                    break
            
            # Add chunk if valid
            if new_chunk:
                chunks.append(' '.join(new_chunk))
                chunk_token_counts.append(new_token_count)
                logger.debug(f"Chunk {len(chunks)} created: {new_token_count} tokens (overlap: {overlap_token_count})")
            else:
                break
    
    logger.debug(f"Text split into {len(chunks)} chunks with token counts: {chunk_token_counts}")
    logger.debug(f"Sum of chunk tokens: {sum(chunk_token_counts)}")
    return chunks, chunk_token_counts

# Split dialogue into chunks for summarization, preserving speaker segments
def split_dialogue_into_chunks(dialogue, tokenizer, max_tokens=1000, overlap_tokens=300):
    """Split dialogue into chunks with specified token counts and overlaps, preserving speaker segments."""
    chunks = []
    chunk_token_counts = []
    current_chunk = []
    current_token_count = 0
    segment_index = 0
    
    while segment_index < len(dialogue):
        # Create first chunk (~1000 tokens)
        if not chunks:
            while segment_index < len(dialogue) and current_token_count < max_tokens:
                speaker, text, start, end = dialogue[segment_index]
                segment_text = f"{speaker}: {text}"
                segment_tokens = len(tokenizer(segment_text, return_tensors="pt").input_ids[0])
                if current_token_count + segment_tokens <= max_tokens or not current_chunk:
                    current_chunk.append((speaker, text, start, end))
                    current_token_count += segment_tokens
                    segment_index += 1
                else:
                    break
            chunks.append(current_chunk)
            chunk_token_counts.append(current_token_count)
            logger.debug(f"Chunk 1 created: {current_token_count} tokens")
        
        # Create subsequent chunks with overlap (~300 tokens from previous + ~700 new)
        else:
            # Get last ~300 tokens from previous chunk
            prev_chunk = chunks[-1]
            overlap_chunk = []
            overlap_token_count = 0
            for speaker, text, start, end in reversed(prev_chunk):
                segment_text = f"{speaker}: {text}"
                segment_tokens = len(tokenizer(segment_text, return_tensors="pt").input_ids[0])
                if overlap_token_count + segment_tokens <= overlap_tokens or not overlap_chunk:
                    overlap_chunk.insert(0, (speaker, text, start, end))
                    overlap_token_count += segment_tokens
                else:
                    break
            
            # Add new ~700 tokens
            new_chunk = overlap_chunk.copy()
            new_token_count = overlap_token_count
            while segment_index < len(dialogue) and new_token_count < max_tokens:
                speaker, text, start, end = dialogue[segment_index]
                segment_text = f"{speaker}: {text}"
                segment_tokens = len(tokenizer(segment_text, return_tensors="pt").input_ids[0])
                if new_token_count + segment_tokens <= max_tokens or not new_chunk:
                    new_chunk.append((speaker, text, start, end))
                    new_token_count += segment_tokens
                    segment_index += 1
                else:
                    break
            
            # Add chunk if valid
            if new_chunk:
                chunks.append(new_chunk)
                chunk_token_counts.append(new_token_count)
                logger.debug(f"Chunk {len(chunks)} created: {new_token_count} tokens (overlap: {overlap_token_count})")
            else:
                break
    
    logger.debug(f"Dialogue split into {len(chunks)} chunks with token counts: {chunk_token_counts}")
    logger.debug(f"Sum of chunk tokens: {sum(chunk_token_counts)}")
    return chunks, chunk_token_counts

# Process audio file for transcription and optional summarization
def process_audio(audio_file, transcription_type, language, model_scale, include_summary):
    try:
        # Verify ffmpeg is installed
        subprocess.run(['ffmpeg', '-version'], check=True, capture_output=True)
    except FileNotFoundError:
        raise Exception("ffmpeg is not installed or not found in the system PATH")

    # Validate audio file
    if not audio_file or not audio_file.filename:
        raise Exception("No valid audio file provided")
    
    start_time = time.time()
    # Generate unique filename and save audio
    audio_filename = f"audio_{uuid.uuid4()}{os.path.splitext(audio_file.filename)[1]}"
    audio_path = os.path.join(app.config['UPLOAD_FOLDER'], audio_filename)
    
    # Save the file and verify it exists
    audio_file.save(audio_path)
    if not os.path.isfile(audio_path):
        raise Exception(f"Failed to save audio file: {audio_path}")

    # Initialize results dictionary
    device = "cpu"
    batch_size = 4
    raw_datetime = datetime.datetime.now(pytz.timezone('Asia/Ho_Chi_Minh')).strftime("%d/%m/%Y %H:%M:%S UTC%z")
    formatted_datetime = f"{raw_datetime[:-2]}:{raw_datetime[-2:]}"
    results = {
        'basic_transcription': '',
        'transcription_with_speakers': [],
        'basic_summary': '',
        'basic_summary_paragraphs': [],
        'dialogue_summary': '',
        'dialogue_summary_paragraphs': [],
        'basic_chunks': [],
        'basic_chunk_token_counts': [],
        'basic_chunk_summaries': [],
        'dialogue_chunks': [],
        'dialogue_chunk_token_counts': [],
        'dialogue_chunk_summaries': [],
        'processing_time': 0,
        'formatted_processing_time': '',
        'date_time': formatted_datetime,
        'transcription_type': transcription_type,
        'language': language,
        'model_scale': model_scale,
        'token_count': 0,
        'number_of_speakers': 0,
        'audio_filename': audio_file.filename,
        'audio_path': audio_path
    }

    # Load and transcribe audio using Whisper model
    model_name = model_scale if language == "english-only" else model_scale
    whisper_model = whisper.load_model(model_name, device=device)
    whisper_result = whisper_model.transcribe(audio_path)
    
    # Process basic transcription (no speakers)
    if transcription_type == "without_speakers":
        basic_text = whisper_result["text"].strip()
        sentences = re.split(r'(?<=[.!?])\s+', basic_text)
        results['basic_transcription'] = " ".join(sentence.strip() for sentence in sentences if sentence.strip())
    else:
        results['basic_transcription'] = whisper_result["text"].strip()

    # Initialize tokenizer for token counting
    tokenizer = AutoTokenizer.from_pretrained("knkarthick/MEETING_SUMMARY")

    # Calculate token count for basic transcription
    basic_text = results['basic_transcription']
    tokens = []
    for i in range(0, len(basic_text), 1000):
        chunk = basic_text[i:i+1000]
        chunk_tokens = tokenizer(chunk, return_tensors="pt", truncation=False).input_ids[0]
        tokens.extend(chunk_tokens)
    results['token_count'] = len(tokens)
    logger.debug(f"Basic transcription token count: {results['token_count']}")

    # Process transcription with speaker diarization
    if transcription_type == "with_speakers":
        # Load audio and WhisperX model
        audio = whisperx.load_audio(audio_path)
        whisperx_model = whisperx.load_model(model_name, device, compute_type="float32")
        whisperx_result = whisperx_model.transcribe(audio, batch_size=batch_size)
        
        # Align transcription with timestamps
        align_model, metadata = whisperx.load_align_model(language_code=whisperx_result["language"], device=device)
        aligned_result = whisperx.align(whisperx_result["segments"], align_model, metadata, audio, device)
        
        # Perform speaker diarization
        diarize_model = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token=HF_TOKEN)
        diarize_segments = diarize_model(audio_path)
        
        # Convert diarization results to DataFrame
        segments = [{'start': seg.start, 'end': seg.end, 'speaker': label} for seg, _, label in diarize_segments.itertracks(yield_label=True)]
        diarize_df = pd.DataFrame(segments)
        
        # Assign speakers to transcription segments
        result_with_speakers = whisperx.assign_word_speakers(diarize_df, aligned_result)
        
        # Remap speaker labels (e.g., SPEAKER A, B, etc.)
        speaker_mapping = {}
        next_id = 0
        dialogue = []
        current_speaker = None
        current_text = ""
        current_start = None
        for segment in result_with_speakers["segments"]:
            original_speaker = segment.get("speaker", "Unknown Speaker")
            if original_speaker not in speaker_mapping:
                speaker_letter = chr(65 + next_id)
                speaker_mapping[original_speaker] = f"SPEAKER {speaker_letter}"
                next_id += 1
            speaker = speaker_mapping[original_speaker]
            text = segment["text"].strip()
            start = segment["start"]
            end = segment["end"]
            if speaker == current_speaker:
                current_text += " " + text
            else:
                if current_speaker is not None:
                    dialogue.append((current_speaker, current_text.strip(), current_start, end))
                current_speaker = speaker
                current_text = text
                current_start = start
        if current_speaker is not None:
            dialogue.append((current_speaker, current_text.strip(), current_start, end))
        
        results['transcription_with_speakers'] = [(speaker, text, start, end) for speaker, text, start, end in dialogue]
        results['number_of_speakers'] = len(set(speaker for speaker, _, _, _ in dialogue))

        # Calculate token count for dialogue
        dialogue_text = "\n".join([f"{speaker}: {text}" for speaker, text, _, _ in dialogue])
        tokens = []
        for i in range(0, len(dialogue_text), 1000):
            chunk = dialogue_text[i:i+1000]
            chunk_tokens = tokenizer(chunk, return_tensors="pt", truncation=False).input_ids[0]
            tokens.extend(chunk_tokens)
        results['token_count'] = len(tokens)
        logger.debug(f"Dialogue transcription token count: {results['token_count']}")

    # Generate summaries if requested
    if include_summary:
        summarizer = pipeline("summarization", model="knkarthick/MEETING_SUMMARY")
        
        # Summarize text or dialogue, handling chunking
        def summarize_text(text, max_chunk_tokens=1000, is_dialogue=False, dialogue=None):
            if not text.strip():
                logger.warning("Empty text provided to summarize_text")
                return "", [], [], [], []
            
            total_tokens = 0
            tokens = []
            for i in range(0, len(text), 1000):
                chunk = text[i:i+1000]
                chunk_tokens = tokenizer(chunk, return_tensors="pt", truncation=False).input_ids[0]
                tokens.extend(chunk_tokens)
            total_tokens = len(tokens)
            max_length = min(math.ceil(total_tokens / 2), 512)
            min_length = min(math.ceil(max_length / 2), 256)
            logger.debug(f"Total tokens: {total_tokens}, max_length: {max_length}, min_length: {min_length}")

            if total_tokens < 1000:
                try:
                    input_length = len(tokenizer(text, return_tensors="pt", truncation=True, max_length=1024).input_ids[0])
                    adjusted_max_length = min(max_length, max(input_length // 2, 30))
                    adjusted_min_length = min(min_length, max(adjusted_max_length // 2, 10))
                    summary = summarizer(text, max_length=adjusted_max_length, min_length=adjusted_min_length, do_sample=False)[0]['summary_text']
                    logger.debug(f"Direct summary generated: {summary[:100]}...")
                    return summary, [], [], [], [summary]
                except Exception as e:
                    logger.error(f"Error in direct summarization: {str(e)}")
                    return f"Error summarizing text: {str(e)}", [], [], [], []
            
            try:
                if is_dialogue:
                    chunks, chunk_token_counts = split_dialogue_into_chunks(dialogue, tokenizer, max_tokens=max_chunk_tokens)
                    chunk_texts = ["\n".join([f"{speaker}: {text}" for speaker, text, _, _ in chunk]) for chunk in chunks]
                else:
                    chunks, chunk_token_counts = split_text_into_chunks(text, tokenizer, max_tokens=max_chunk_tokens)
                    chunk_texts = chunks
                
                logger.debug(f"Number of chunks: {len(chunks)}")
                if not chunks:
                    logger.warning("No valid chunks generated")
                    return "No valid chunks to summarize", [], [], [], []
                
                chunk_summaries = []
                for i, chunk_text in enumerate(chunk_texts):
                    try:
                        input_length = len(tokenizer(chunk_text, return_tensors="pt", truncation=True, max_length=1024).input_ids[0])
                        adjusted_max_length = min(max_length, max(input_length // 2, 30))
                        adjusted_min_length = min(min_length, max(adjusted_max_length // 2, 10))
                        summary = summarizer(chunk_text, max_length=adjusted_max_length, min_length=adjusted_min_length, do_sample=False)[0]['summary_text']
                        chunk_summaries.append(summary)
                        logger.debug(f"Chunk {i+1} summary: {summary[:100]}...")
                    except Exception as e:
                        logger.error(f"Error summarizing chunk {i+1}: {str(e)}")
                        chunk_summaries.append(f"Error summarizing chunk: {str(e)}")
                
                combined_summaries = '\n\n'.join(chunk_summaries)
                if not combined_summaries.strip():
                    logger.warning("Combined summaries are empty")
                    return "No valid summaries generated", [], [], [], []
                
                final_summary = combined_summaries
                logger.debug(f"Final summary (combined chunks): {final_summary[:100]}...")
                return final_summary, chunks, chunk_token_counts, chunk_summaries, chunk_summaries
            except Exception as e:
                logger.error(f"Error in summarization: {str(e)}")
                return f"Error in summarization: {str(e)}", chunks, chunk_token_counts, chunk_summaries, []

        # Summarize basic transcription
        results['basic_summary'], results['basic_chunks'], results['basic_chunk_token_counts'], results['basic_chunk_summaries'], results['basic_summary_paragraphs'] = summarize_text(
            basic_text
        )
        
        # Summarize dialogue if speakers are included
        if transcription_type == "with_speakers":
            dialogue_text = "\n".join([f"{speaker}: {text}" for speaker, text, _, _ in results['transcription_with_speakers']])
            results['dialogue_summary'], results['dialogue_chunks'], results['dialogue_chunk_token_counts'], results['dialogue_chunk_summaries'], results['dialogue_summary_paragraphs'] = summarize_text(
                dialogue_text, is_dialogue=True, dialogue=results['transcription_with_speakers']
            )

    # Calculate processing time
    results['processing_time'] = time.time() - start_time
    results['formatted_processing_time'] = format_processing_time(results['processing_time'])

    # Generate DOCX report with transcription and summaries
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    try:
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(12)
    except Exception as e:
        logger.warning(f"Failed to set Times New Roman font: {str(e)}. Using default font.")

    # Add title to document
    p = doc.add_paragraph()
    run = p.add_run("TRANSCRIPTION DETAILS")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Add metadata to document
    for label, value in [
        ("Audio:", results['audio_filename']),
        ("Date and Time:", results['date_time']),
        ("Processing Time:", results['formatted_processing_time']),
        ("Transcription Type:", transcription_type.replace('_', ' ').title()),
        ("Language:", language.replace('-', ' ').title()),
        ("Model Scale:", model_scale),
        ("Total Tokens:", str(results['token_count'])),
    ] + ([("Number of Speakers:", str(results['number_of_speakers']))] if transcription_type == "with_speakers" else []):
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        run = p.add_run(f" {value}")
        run.bold = False
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    doc.add_paragraph()

    # Add transcription without speakers to document
    if transcription_type == "without_speakers":
        p = doc.add_paragraph()
        run = p.add_run("BASIC TRANSCRIPTION")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run(basic_text)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run.font.size = Pt(12)
        
        if include_summary:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("SUMMARY OF BASIC TRANSCRIPTION")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            for summary in results['basic_summary_paragraphs']:
                p = doc.add_paragraph()
                run = p.add_run(summary)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph()
            
            if results['basic_chunks']:
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("CHUNK DETAILS")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                
                for i, (chunk, chunk_summary) in enumerate(zip(results['basic_chunks'], results['basic_chunk_summaries'])):
                    p = doc.add_paragraph()
                    run = p.add_run(f"Chunk {i+1}: {results['basic_chunk_token_counts'][i]} tokens")
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    p = doc.add_paragraph()
                    run = p.add_run(chunk)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    run = p.add_run(f"Summary of Chunk {i+1}")
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    p = doc.add_paragraph()
                    run = p.add_run(chunk_summary)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    doc.add_paragraph()
    # Add transcription with speakers to document
    else:
        p = doc.add_paragraph()
        run = p.add_run("TRANSCRIPTION WITH SPEAKERS")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        for speaker, text, _, _ in results['transcription_with_speakers']:
            p = doc.add_paragraph()
            run = p.add_run(speaker)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph()
        
        if include_summary:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("SUMMARY OF TRANSCRIPTION WITH SPEAKERS")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            for summary in results['dialogue_summary_paragraphs']:
                p = doc.add_paragraph()
                run = p.add_run(summary)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph()
            
            if results['dialogue_chunks']:
                doc.add_paragraph()
                p = doc.add_paragraph()
                run = p.add_run("CHUNK DETAILS")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                
                for i, chunk in enumerate(results['dialogue_chunks']):
                    p = doc.add_paragraph()
                    run = p.add_run(f"Chunk {i+1}: {results['dialogue_chunk_token_counts'][i]} tokens")
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    for speaker, text, _, _ in chunk:
                        p = doc.add_paragraph()
                        run = p.add_run(speaker)
                        run.bold = True
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.bold = False
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        doc.add_paragraph()
                    
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    run = p.add_run(f"Summary of Chunk {i+1}")
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    
                    p = doc.add_paragraph()
                    run = p.add_run(results['dialogue_chunk_summaries'][i])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    doc.add_paragraph()

    # Save DOCX file and update results
    doc_file = f"transcription_{uuid.uuid4()}.docx"
    doc_path = os.path.join(OUTPUT_FOLDER, doc_file)
    doc.save(doc_path)

    results['doc_path'] = doc_path
    return results

# Route to serve the main page
@app.route('/')
def index():
    return render_template('index.html')

# Route to handle audio file uploads and processing
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'audio_file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    audio_file = request.files['audio_file']
    if not audio_file.filename:
        return jsonify({'error': 'No valid audio file selected'}), 400

    transcription_type = request.form.get('transcription_type', 'without_speakers')
    language = request.form.get('language', 'english-only')
    model_scale = request.form.get('model_scale', 'base.en' if language == 'english-only' else 'base')
    include_summary = request.form.get('include_summary', 'false') == 'true'

    try:
        results = process_audio(audio_file, transcription_type, language, model_scale, include_summary)
        return jsonify(results)
    except Exception as e:
        logger.error(f"Error in upload: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Route to serve the generated DOCX file for download
@app.route('/download/<path:filename>')
def download_file(filename):
    return send_file(os.path.join(OUTPUT_FOLDER, filename), as_attachment=True)

# Route to serve the uploaded audio file
@app.route('/audio/<path:filename>')
def serve_audio(filename):
    return send_file(os.path.join(UPLOAD_FOLDER, filename))

# Run the Flask app
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)